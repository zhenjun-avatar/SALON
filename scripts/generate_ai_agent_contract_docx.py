"""One-off generator: AI agent service contract Word doc (simplified fee schemes)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent.parent / "assets" / "ai-agent-service-contract.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")
    return p


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_table_fee_package(doc: Document):
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    hdr = ("档位", "适用场景", "包含内容（示例）", "含税总价（人民币）", "备注")
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        (
            "标准版",
            "多工具 + 知识库 + 1 套系统集成",
            "工作流 + RAG（单库）、企微/飞书等 1 类集成、观测与基础评测、文档与 2 次培训",
            "¥_______ 万",
            "周期约 6–10 周",
        ),
        (
            "增强版",
            "多场景 / 混合检索 / 多系统",
            "多工作流协调、混合检索与重排、多数据源、2 类及以上系统对接、联调与压测支持",
            "¥_______ 万",
            "周期双方另议",
        ),
    ]
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = cell
    doc.add_paragraph()


def add_table_fee_build_maint(doc: Document):
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = ("项目", "金额（人民币）", "说明")
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    body = [
        ("首期建设费", "¥_______", "对应「按项目打包」所选档位（标准版 / 增强版），或双方书面确认的金额"),
        ("年度运维与支持", "¥_______ / 年", "约定 SLA 内响应、小版本修复、每月 ______ 小时优化额度；不含大功能新增"),
        ("超出额度", "¥_______ / 小时 或 人天计价", "须事先书面确认"),
    ]
    for ri, row in enumerate(body, start=1):
        for ci, text in enumerate(row):
            table.rows[ri].cells[ci].text = text
    doc.add_paragraph()


def main():
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = sect.bottom_margin = Cm(2.5)
    sect.left_margin = sect.right_margin = Cm(2.5)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI 代理构建服务合同")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "黑体")

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("（草案 · 供磋商使用）")
    sr.font.size = Pt(11)
    sr.font.name = "宋体"
    sr._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")

    add_para(doc, "合同编号：_______________　　签订日期：_______________　　签订地点：_______________")

    add_heading(doc, "甲方（委托方）", level=2)
    add_para(doc, "名称：________________________________________________")
    add_para(doc, "统一社会信用代码 / 身份证号：________________________________________________")
    add_para(doc, "地址：________________________________________________")
    add_para(doc, "联系人：____________　职务：____________　电话：____________　电子邮箱：____________")

    add_heading(doc, "乙方（服务方）", level=2)
    add_para(doc, "名称：________________________________________________")
    add_para(doc, "统一社会信用代码 / 身份证号：________________________________________________")
    add_para(doc, "地址：________________________________________________")
    add_para(doc, "联系人：____________　职务：____________　电话：____________　电子邮箱：____________")

    add_para(
        doc,
        "甲、乙双方本着平等自愿、诚实信用原则，就甲方委托乙方提供企业级 AI 代理（Agent）构建与相关技术服务事宜，达成如下条款。",
    )

    add_heading(doc, "第一条　服务内容与范围", level=2)
    add_para(
        doc,
        "乙方根据甲方需求，在约定范围内提供下列一项或多项服务（以《工作说明书》（SOW）或附件清单为准）：",
    )
    for item in (
        "需求与方案：业务场景梳理、人机协作边界、数据与系统集成清单、技术路线说明。",
        "智能体与工作流：基于 Dify、LangGraph 等工具的可视化编排、工具调用（HTTP/API）、条件分支、会话与状态设计等。",
        "知识库与 RAG：文档解析与切分策略、向量与检索方案、元数据与权限设计；可选混合检索、重排序、评测与观测。",
        "系统集成：与甲方现有系统（如企微、飞书多维表、业务网关、数据库等）的对接与联调支持。",
        "交付与移交：可运行环境说明、配置与密钥管理建议、基础运维与变更说明；约定次数内的培训与答疑。",
    ):
        doc.add_paragraph(item, style="List Number")
    add_para(
        doc,
        "不在默认范围内（除非 SOW 另行约定并计费）：甲方自有模型采购费、第三方 SaaS/API 按量费用、超出约定范围的定制开发、7×24 驻场运维、等保/密评等合规认证代办。",
    )

    add_heading(doc, "第二条　交付物与验收", level=2)
    add_para(doc, "交付物以 SOW 列明的软件制品、配置文件、文档、演示环境或仓库访问权限为准。")
    add_para(
        doc,
        "甲方应在收到乙方验收申请后 ______ 个工作日内完成验收；逾期未书面提出具体异议的，视为阶段性验收通过（双方另有约定的除外）。",
    )
    add_para(
        doc,
        "验收标准：符合 SOW 中的功能描述与测试用例；重大缺陷（导致约定核心流程不可用）未修复的，不视为通过。",
    )

    add_heading(doc, "第三条　费用方案", level=2)
    add_para(doc, "本合同费用按下列方案之一执行，由双方在《报价单》中勾选确认（可同时约定首期 + 年度运维的组合）。")

    add_heading(doc, "方案一：按项目打包（固定价）", level=3)
    add_para(doc, "以下两档为固定总价包干（含税），以双方盖章《报价单》所载金额为准。")
    add_table_fee_package(doc)
    add_para(doc, "说明：上表总价包含第三条方案约定范围内的开发、联调、文档及约定培训；不含第三方 API/SaaS 按量费用。")

    add_heading(doc, "方案二：首期建设 + 年度运维", level=3)
    add_para(
        doc,
        "首期建设费与「方案一」所选档位挂钩，或由双方书面约定固定金额；年度运维为订阅制，按年支付。",
    )
    add_table_fee_build_maint(doc)

    add_heading(doc, "第四条　付款方式", level=2)
    for line in (
        "签约后 ______ 日内，甲方向乙方支付合同总额的 ______ % 作为预付款。",
        "里程碑付款（可勾选）：需求确认 ______ %；核心联调通过 ______ %；终验通过 ______ %；质保期满 ______ %。",
        "发票类型：____________；甲方收到合格发票后 ______ 日内付款。",
        "乙方收款账户：户名 ____________；开户行 ____________；账号 ____________。",
    ):
        add_para(doc, line)

    add_heading(doc, "第五条　双方权利义务", level=2)
    add_para(
        doc,
        "甲方应及时提供业务资料、测试账号、接口文档与决策人，并对数据合法来源与授权负责。",
    )
    add_para(
        doc,
        "乙方应按约定专业标准交付，对已知安全风险与架构限制予以说明；未经甲方同意不得将甲方数据用于与本项目无关用途。",
    )

    add_heading(doc, "第六条　知识产权", level=2)
    add_para(
        doc,
        "在甲方付清本合同项下应付费用后，SOW 约定为甲方专有的定制交付物（不含乙方通用框架、开源组件），甲方取得约定范围内的使用权；著作权归属可由附件另行约定。乙方自有可复用资产，乙方保留知识产权，甲方获得本项目范围内使用权。",
    )

    add_heading(doc, "第七条　保密", level=2)
    add_para(doc, "双方对履行合同中知悉的对方商业秘密、技术信息、用户数据等承担保密义务，期限为合同终止后 ______ 年。")

    add_heading(doc, "第八条　质保与维护", level=2)
    add_para(
        doc,
        "免费质保期：终验通过之日起 ______ 个月，修复 SOW 范围内缺陷（非新需求）。质保期外维护按方案二或另行报价执行。",
    )

    add_heading(doc, "第九条　违约责任", level=2)
    add_para(
        doc,
        "甲方逾期付款，每逾期一日按逾期金额的万分之 ______ 计违约金（上限不超过逾期部分 ______ %）。",
    )
    add_para(
        doc,
        "乙方无正当理由逾期交付超过 ______ 日，甲方可催告；仍不改正的，可解除合同并就未履行部分主张退款（比例可附件约定）。因不可抗力或甲方原因导致延期，工期顺延。",
    )

    add_heading(doc, "第十条　合同变更与解除", level=2)
    add_para(
        doc,
        "范围、费用、周期变更须书面补充协议或变更单。一方严重违约且未在 ______ 日内补救的，守约方可解除合同并主张直接损失（间接损失可约定限额或排除）。",
    )

    add_heading(doc, "第十一条　争议解决", level=2)
    add_para(
        doc,
        "协商不成的，提交 ________________________ 人民法院诉讼解决（或 ________________________ 仲裁委员会仲裁）。",
    )

    add_heading(doc, "第十二条　其他", level=2)
    add_para(doc, "附件与正文具有同等效力：《工作说明书（SOW）》《报价单》《保密协议》（如有）。")
    add_para(doc, "本合同一式 ______ 份，双方各执 ______ 份，自双方盖章之日起生效。")

    doc.add_paragraph()
    add_para(doc, "甲方（盖章）：____________________　　授权代表：____________________　　日期：__________")
    add_para(doc, "乙方（盖章）：____________________　　授权代表：____________________　　日期：__________")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
